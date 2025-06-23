import io
import re

from docx import Document


def apply_speaker_names_to_text(text: str, speaker_name_mapping: dict) -> str:
    """
    將發言者名稱映射應用到文字內容中
    
    Args:
        text (str): 原始文字內容
        speaker_name_mapping (dict): 發言者名稱映射
    
    Returns:
        str: 應用發言者名稱映射後的文字內容
    """
    if not speaker_name_mapping or not text:
        return text
    
    modified_text = text
    for original_speaker, custom_name in speaker_name_mapping.items():
        # 替換各種可能的發言者格式
        patterns = [
            rf'{re.escape(original_speaker)}：',  # 發言者00：
            rf'{re.escape(original_speaker)}:',   # 發言者00:
            rf'\b{re.escape(original_speaker)}\b',  # 單獨的發言者名稱
        ]
        
        for pattern in patterns:
            modified_text = re.sub(pattern, custom_name, modified_text)
    
    return modified_text


def format_summary_for_export(summary_data: dict, meeting_info, speaker_name_mapping: dict = None) -> str:
    """
    將摘要資料格式化為純文字字串。
    
    Args:
        summary_data (dict): 摘要資料
        meeting_info: 會議資訊
        speaker_name_mapping (dict): 發言者名稱映射
    """
    lines = []
    # 安全地獲取會議檔名
    original_filename = meeting_info['original_filename'] if meeting_info and 'original_filename' in meeting_info else 'N/A'
    lines.append(f"會議摘要：{original_filename}")
    lines.append("=" * 40)
    lines.append(f"生成時間：{summary_data.get('summary_generated_at', 'N/A')}")
    lines.append("\n")

    lines.append("【會議總結】")
    lines.append("-" * 20)
    global_summary = summary_data.get('global_summary', '無內容')
    if speaker_name_mapping:
        global_summary = apply_speaker_names_to_text(global_summary, speaker_name_mapping)
    lines.append(global_summary)
    lines.append("\n")

    lines.append("【發言人重點】")
    lines.append("-" * 20)
    highlights = summary_data.get('speaker_highlights', [])
    if highlights:
        for item in highlights:
            if speaker_name_mapping:
                item = apply_speaker_names_to_text(item, speaker_name_mapping)
            lines.append(f"• {item}")
    else:
        lines.append("無內容")
    lines.append("\n")

    lines.append("【分段摘要】")
    lines.append("-" * 20)
    chunks = summary_data.get('chunk_summaries', '')
    if chunks:
        if speaker_name_mapping:
            chunks = apply_speaker_names_to_text(chunks, speaker_name_mapping)
        # 假設分段摘要是以換行符分隔的
        chunk_list = chunks.strip().split('\n')
        for i, chunk in enumerate(chunk_list):
            lines.append(f"片段 {i+1}: {chunk.lstrip('- ')}")
    else:
        lines.append("無內容")

    return "\n".join(lines)


def create_summary_docx(summary_data: dict, meeting_info, speaker_name_mapping: dict = None) -> io.BytesIO:
    """
    從摘要資料創建 Word 文件。
    
    Args:
        summary_data (dict): 摘要資料
        meeting_info: 會議資訊
        speaker_name_mapping (dict): 發言者名稱映射
    """
    document = Document()
    
    # 安全地獲取會議檔名
    original_filename = meeting_info['original_filename'] if meeting_info and 'original_filename' in meeting_info else 'N/A'
    
    # 設置文件標題
    document.add_heading(f"會議摘要：{original_filename}", level=0)
    
    # 添加生成時間
    p = document.add_paragraph()
    p.add_run('生成時間：').italic = True
    p.add_run(str(summary_data.get('summary_generated_at', 'N/A')))
    
    # 添加總結
    document.add_heading('會議總結', level=1)
    global_summary = summary_data.get('global_summary', '無內容')
    if speaker_name_mapping:
        global_summary = apply_speaker_names_to_text(global_summary, speaker_name_mapping)
    document.add_paragraph(global_summary)

    # 添加發言人重點
    document.add_heading('發言人重點', level=1)
    highlights = summary_data.get('speaker_highlights', [])
    if highlights:
        for item in highlights:
            if speaker_name_mapping:
                item = apply_speaker_names_to_text(item, speaker_name_mapping)
            document.add_paragraph(item, style='List Bullet')
    else:
        document.add_paragraph('無內容')

    # 添加分段摘要
    document.add_heading('分段摘要', level=1)
    chunks = summary_data.get('chunk_summaries', '')
    if chunks:
        if speaker_name_mapping:
            chunks = apply_speaker_names_to_text(chunks, speaker_name_mapping)
        chunk_list = chunks.strip().split('\n')
        for i, chunk in enumerate(chunk_list):
            p = document.add_paragraph()
            p.add_run(f'片段 {i+1}: ').bold = True
            p.add_run(chunk.lstrip('- '))
    else:
        document.add_paragraph('無內容')

    # 將文件保存在內存中
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    return file_stream 